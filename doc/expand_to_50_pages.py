from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path("Gojo_Sentinel_APA7_Chapters_1_to_5.docx")
OUTPUT = Path("Gojo_Sentinel_APA7_Expanded_50_Pages.docx")


def set_run_font(run):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(14)


def normalize(doc):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        for p in section.header.paragraphs:
            p.text = ""
        for p in section.footer.paragraphs:
            p.text = ""

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_after = Pt(0)

    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        if p.style.name == "Normal" and p.text.strip() and not p.text.strip().startswith("[Insert"):
            p.paragraph_format.first_line_indent = Inches(0.5)
        for run in p.runs:
            set_run_font(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.first_line_indent = None
                    for run in p.runs:
                        set_run_font(run)


def add_para(doc, text, style="Normal", align=None, bold=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    if style == "Normal":
        p.paragraph_format.first_line_indent = Inches(0.5)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run)
    return p


def add_heading(doc, text, level=2):
    p = add_para(doc, text, style=f"Heading {level}", bold=True)
    p.paragraph_format.first_line_indent = None
    return p


EXPANSIONS = {
    "CHAPTER ONE": [
        ("1.11 Background of the Application Domain", [
            "The application domain of this project is digital financial security. In modern banking, the transaction environment is no longer limited to physical branches or teller counters. Customers now interact with banks through mobile applications, web portals, USSD codes, POS terminals, and third-party fintech platforms. This shift has created a large digital transaction space where speed, convenience, and availability are major advantages. However, the same digital channels also increase the number of points through which fraud can occur. A transaction may be initiated from a mobile phone, routed through an inter-bank payment platform, and received by a bank or wallet account within a very short time. This speed is useful to customers but challenging for fraud investigators because suspicious activity must be identified before the money is moved beyond recovery (CBN, 2024; NIBSS, 2024).",
            "In the Nigerian financial environment, digital transactions have become part of daily life for students, traders, salary earners, business owners, and government workers. Many people rely on instant transfers and mobile banking for basic payments. As a result, fraud detection systems must understand local payment behavior. A model designed only with foreign transaction patterns may not fully capture the behavior of Nigerian users, where USSD banking, fintech wallets, NIP transfers, and BVN verification are common. The proposed system therefore focuses on transaction features that are realistic within the Nigerian context, including sender bank, receiver bank, payment channel, BVN match status, transaction amount, and transaction time.",
            "The need for intelligent fraud detection is also connected to customer trust. When customers experience unauthorized transactions, they may lose confidence in digital banking services. Financial institutions therefore need systems that can act quickly, provide reasonable explanations, and support human administrators with evidence. Gojo Sentinel was designed with this concern in mind. It does not only give a fraud score; it also presents risk indicators so that the user can understand why a transaction appears safe or suspicious.",
        ]),
        ("1.12 Motivation for the Study", [
            "The motivation for this study came from the increasing dependence on electronic payments and the growing concern about digital banking fraud. While digital banking has improved access to financial services, many users still face the risk of unauthorized transfers, phishing-related attacks, identity misuse, and suspicious account activity. In many cases, fraud is noticed only after the transaction has been completed. This creates a problem because digital money can be transferred across several accounts within minutes.",
            "Another motivation is the limitation of relying only on fixed rules. A rule that blocks all transactions above a certain amount may be useful, but it may also stop genuine high-value transactions. At the same time, a fraudster may perform several smaller transactions to avoid the rule. This shows that fraud detection should consider more than one factor. A machine learning model can study different transaction features together and produce a probability score rather than depending only on one threshold (Hernandez Aros et al., 2024; Naqvi, 2025).",
            "The project also has academic motivation. It gives a practical example of how software development and artificial intelligence can be combined to solve a security problem. The work involves data generation, preprocessing, model training, API development, frontend design, database management, and testing. These activities show the relationship between theory and practice in computer science.",
        ]),
        ("1.13 Justification of the Study", [
            "The study is justified because fraud detection remains a major requirement in digital banking. A system that can score transactions in real time can help financial institutions respond faster to suspicious activity. Although the project is a prototype, it demonstrates an approach that can be improved and adapted for real banking environments. It also shows that fraud detection should not be treated only as a manual process but as a data-driven decision-support system.",
            "The use of machine learning is justified because fraud patterns are not always simple. A transaction may be risky because of the combination of amount, channel, time, receiver type, and BVN status. A human reviewer may miss these patterns when many transactions are being processed. Machine learning models can process several features at the same time and produce risk scores that support faster decisions (Detthamrong et al., 2024; Nobel et al., 2024).",
            "The Nigerian focus of the system is also important. Many payment environments have local characteristics. For example, USSD is widely used because it works without internet access, while fintech wallets are common for fast transactions. BVN verification is also a major identity feature in Nigerian banking. Designing the system around these features makes the project more relevant to the intended environment.",
        ]),
        ("1.14 Expected Benefits of the System", [
            "The expected benefit of Gojo Sentinel is that it can support faster detection of suspicious transactions. Instead of waiting for manual review after a complaint has been made, the system can provide an immediate risk score. This can help administrators decide whether a transaction should be approved, reviewed, declined, or blocked.",
            "Another benefit is improved monitoring. Since the system stores prediction history, administrators can review previous transactions and observe patterns over time. This can support internal reporting, investigation, and future model improvement. The dashboard also makes the system easier to use because results are presented visually instead of requiring users to read raw API responses.",
            "The system can also serve as a foundation for future research. Other students can improve the model, add real data, connect live APIs, strengthen security, or compare different algorithms. In this way, the project is not only a final-year implementation but also a base for continued development.",
        ]),
    ],
    "CHAPTER TWO: LITERATURE REVIEW": [
        ("2.16 Theoretical Foundation of Fraud Detection", [
            "Fraud detection is built on the idea that fraudulent behavior usually differs from normal behavior in one or more measurable ways. These differences may be obvious, such as a very large transfer, or subtle, such as a transaction made at an unusual time through a channel the customer rarely uses. The main task of a fraud detection system is to identify these differences early enough to prevent or reduce loss. In machine learning, this is treated as a classification problem where the model learns from examples and predicts whether a new transaction is likely to be legitimate or fraudulent.",
            "The theoretical foundation of this project is connected to pattern recognition. Pattern recognition assumes that data contains signals that can be used to distinguish one class from another. In the case of financial transactions, the signal may come from amount, time, user activity, payment channel, or identity information. A machine learning algorithm searches for relationships between these features and the final fraud label. Once trained, the model applies what it has learned to new transactions (Hernandez Aros et al., 2024; Sensors, 2024).",
            "Another important concept is risk scoring. A good fraud detection system should not always make a hard decision immediately. In some cases, it is better to assign a score and allow the organization to decide what level of risk is acceptable. This is why the proposed system returns both a fraud probability and a recommendation. A low-risk transaction can be approved, a medium-risk transaction can be reviewed, and a critical-risk transaction can be blocked.",
        ]),
        ("2.17 Artificial Intelligence and Financial Security", [
            "Artificial intelligence has become important in financial security because modern financial systems produce large volumes of data. Manual monitoring alone cannot keep up with the speed and number of transactions that occur daily. AI-based systems can help by analyzing data quickly and consistently. They can also support security teams by highlighting transactions that deserve attention.",
            "In financial fraud detection, AI is useful because it can identify relationships that are not easy to express as simple rules. For example, a transaction may be suspicious not only because the amount is high but because the high amount is combined with a BVN mismatch, a late-night timestamp, a risky channel, and an unusual receiver bank. A machine learning model can learn such combinations from data and use them to predict risk (Detthamrong et al., 2024; Nobel et al., 2024).",
            "However, AI does not remove the need for human judgment. In real financial institutions, machine learning results are often used as decision support. Human administrators may still review cases, adjust rules, investigate customers, and confirm fraud labels. For this reason, the proposed system includes a dashboard, transaction history, and rules management feature instead of only a prediction model.",
        ]),
        ("2.18 Comparison Between Rule-Based and AI-Based Detection", [
            "Rule-based detection and AI-based detection are both useful, but they solve the problem in different ways. Rule-based systems depend on conditions written by administrators. For instance, a rule may say that any transaction above NGN 2,000,000 should be blocked. This approach is easy to understand and control. It is also useful when an institution has strict policies that must always be applied.",
            "The weakness of rule-based detection is that it does not learn by itself. If fraudsters change their strategy, the rules may no longer be effective. Rule-based systems can also create many false positives when the rules are too strict. For example, a business customer may legitimately transfer a large amount, but the transaction may still be blocked by a simple amount rule.",
            "AI-based detection uses data to learn patterns. Instead of depending only on fixed instructions, the model studies past transactions and estimates the risk of new ones. This makes it more flexible, especially when fraud patterns involve several features at once. The best practical approach is often a hybrid system that combines machine learning scores with rule-based controls (Naqvi, 2025; Preciado Martinez et al., 2025).",
        ]),
        ("2.19 Review of Machine Learning Algorithms for Fraud Detection", [
            "Several machine learning algorithms can be used for fraud detection. Logistic regression is simple and easy to interpret, but it may not capture complex relationships in transaction data. Decision trees are easier to understand because they split data using conditions, but a single tree may overfit the training data. Random Forest improves on decision trees by combining many trees, which usually improves accuracy and stability.",
            "Support Vector Machines can perform well in classification tasks, but they may become slow with large datasets. Neural networks can learn complex patterns, especially when large datasets are available, but they may require more computing resources and may be harder to explain. XGBoost is widely used for tabular data because it combines many weak decision trees into a strong model and often performs well with structured features (Detthamrong et al., 2024; Nobel et al., 2024).",
            "For this project, XGBoost was selected because banking transaction records are structured. The model can work effectively with numerical features such as amount and transaction count, as well as encoded categorical features such as channel and bank name. It is also efficient enough for real-time prediction when the model has already been trained and loaded into memory.",
        ]),
        ("2.20 Research Gap", [
            "From the reviewed literature, it is clear that machine learning has become important in fraud detection. However, many studies focus on general credit card fraud datasets and do not always reflect Nigerian digital payment behavior. Nigerian banking includes local features such as BVN matching, NIP transfers, USSD banking, fintech wallets, and local bank names. These features make the transaction environment different from many public datasets.",
            "Another gap is the need for systems that combine prediction with usability. A model may be accurate, but if users cannot interact with it easily, it becomes difficult to use in practice. Gojo Sentinel addresses this by providing a dashboard where users can submit transactions, view risk scores, monitor history, and manage fraud rules. The project therefore combines machine learning with a practical software interface.",
            "The project also contributes by showing how a prototype can be built when real bank data is not available. Synthetic data is not a perfect replacement for real data, but it allows the design and testing of a complete system. This is useful in academic settings where privacy and access restrictions make real transaction data difficult to obtain (Kennedy et al., 2024; Wen et al., 2024).",
        ]),
        ("2.21 Summary of Reviewed Literature", [
            "The reviewed literature shows that fraud detection has moved from simple manual checks to more intelligent and automated approaches. Rule-based systems remain useful because they are easy to control and explain, but they are limited when fraud patterns change. Machine learning provides a stronger approach because it can learn from data and identify complex relationships among transaction features.",
            "The literature also shows that fraud detection is affected by class imbalance, data quality, model explainability, and real-time performance. Fraud cases are usually fewer than legitimate cases, so the model must be trained carefully to avoid ignoring the minority class. This is why techniques such as SMOTE and careful evaluation are important in fraud-related machine learning projects.",
            "The proposed system builds on these ideas by combining transaction scoring, dashboard interaction, risk indicators, transaction history, and administrative control. It does not claim to solve every fraud problem, but it demonstrates a practical and locally relevant approach to digital banking fraud detection.",
        ]),
    ],
    "CHAPTER THREE": [
        ("3.21 Detailed Requirement Analysis", [
            "Requirement analysis was carried out to identify what the system should do and how users should interact with it. The most important requirement was real-time transaction scoring. The system needed to accept transaction details, process them through a trained model, and return a result quickly. This requirement influenced the choice of FastAPI for the backend and XGBoost for the prediction model.",
            "Another requirement was explainability at the user interface level. Even if the model produces a fraud probability, the user still needs to understand the major risk indicators. For this reason, the frontend displays risk factors such as high amount, BVN mismatch, risky channel, fintech receiver, and late-night transaction. These indicators help the user interpret the model result in a practical way.",
            "Administrative control was also required. The system needed login authentication, user management, transaction history, and rule management. These features make the system more complete because fraud detection is not only about prediction. Administrators must also monitor activity, control access, and update policies when needed (CBN, 2024; Digital Policy Alert, 2024).",
        ]),
        ("3.22 Functional Requirements", [
            "The functional requirements describe the actions the system must perform. The system must allow users to log in, enter transaction details, submit a transaction for scoring, receive a fraud probability, view risk level and recommendation, and check recent transaction results. Administrators must be able to view users, create users, delete users, view fraud rules, create rules, and delete rules.",
            "The prediction function is the central feature. It must receive valid transaction data and return a structured response containing transaction ID, fraud probability, fraud status, risk level, and recommendation. The system must also save the prediction so that it can be reviewed later. This supports accountability and makes it possible to monitor how many transactions have been scored.",
            "The rules function allows administrators to define policies. Even though the current prototype focuses mainly on AI scoring, the presence of the rules interface makes the system ready for future hybrid enforcement. In a production version, rules can be applied before or after the model prediction to support strict business decisions.",
        ]),
        ("3.23 Non-Functional Requirements", [
            "Non-functional requirements describe the qualities the system should have. The first requirement is speed. Since banking transactions are time-sensitive, the system should return prediction results quickly. This is achieved by loading the model at startup so that prediction does not require retraining.",
            "The second requirement is usability. The dashboard should be easy to understand, even for a user who is not a machine learning expert. The system therefore uses a form for transaction entry, a circular score gauge for visual feedback, and clear labels such as LOW, MEDIUM, HIGH, and CRITICAL. These labels make the result easier to interpret.",
            "The third requirement is maintainability. The code is divided into separate files for data generation, preprocessing, training, retraining, backend API, and frontend interface. This separation makes it easier to improve one part of the system without rewriting everything. The fourth requirement is portability. The system can run locally and can also be deployed through Docker or cloud platforms.",
        ]),
        ("3.24 Detailed Data Design", [
            "The data design was created to represent the information needed for fraud detection. Transaction ID identifies each transaction. User ID connects the transaction to a user. Amount shows the value of the transaction. Sender bank and receiver bank describe the direction of money movement. Channel shows whether the transaction was made through NIP, POS, USSD, or Web.",
            "BVN match status is important because identity mismatch can be a warning sign. Timestamp is also important because some fraud attempts happen during late-night periods when customers and bank staff may be less alert. Velocity features such as transaction count within one hour and total amount within twenty-four hours help the system understand recent user behavior.",
            "The design also includes prediction output data. Fraud probability gives a numerical risk score, while risk level and recommendation translate the score into understandable decisions. This is important because administrators may not want to interpret raw probabilities every time they review a transaction.",
        ]),
        ("3.25 Algorithm Design", [
            "The prediction algorithm follows a clear sequence. First, the transaction details are collected from the frontend. Second, the backend validates the request using the expected schema. Third, the data is converted into a Pandas DataFrame so that it can be processed by the saved preprocessing pipeline. Fourth, the pipeline creates the same features used during training.",
            "After preprocessing, the transformed data is passed to the XGBoost model. The model returns the probability that the transaction belongs to the fraud class. The system then compares the probability with defined risk ranges and assigns a risk level. Finally, the result is stored in the database and returned to the frontend.",
            "This design separates prediction from presentation. The backend is responsible for computation, while the frontend is responsible for displaying the result. This makes the system easier to maintain and allows other clients, such as a mobile app, to use the same backend API.",
        ]),
        ("3.26 Interface Design Considerations", [
            "The interface was designed to make the fraud detection process simple for users. A user should not need to understand machine learning before using the system. For this reason, the dashboard provides clearly labelled input fields, channel buttons, simulation buttons, result cards, and risk badges. The result is shown as a risk score, risk level, recommendation, and list of risk indicators.",
            "The system also includes navigation for different administrative tasks. The dashboard is used for scoring, the transactions page is used for history, the rules page is used for fraud policies, and the users page is used for account management. This separation improves usability because each page has a clear purpose.",
            "The design also considers demonstration needs. Since this is an academic project, the Normal Pattern and High-Risk Pattern buttons help during presentation. They allow the student to quickly show how the system responds to different transaction types without manually entering many values.",
        ]),
    ],
    "CHAPTER FOUR": [
        ("4.22 Implementation of the Prediction Workflow", [
            "The prediction workflow begins when the user submits the transaction form on the dashboard. The frontend gathers all required values from the form fields, including transaction ID, user ID, amount, sender bank, receiver bank, channel, sender account number, receiver account number, BVN match status, timestamp, and velocity fields. These values are converted into a JSON request and sent to the backend prediction endpoint.",
            "On the backend, the request is validated using a defined data model. This helps ensure that required fields are present and that the values are in the expected format. The transaction data is then converted into a one-row DataFrame. This is necessary because the preprocessing pipeline expects tabular input similar to the training data.",
            "The preprocessing pipeline transforms the data by extracting time-based features, encoding categorical fields, scaling numerical fields, and dropping fields that are not useful for prediction. The processed data is then passed to the XGBoost model. The model returns a probability score, and the backend converts this score into a risk level and recommendation. This result is saved in the database and displayed on the dashboard.",
        ]),
        ("4.23 Implementation of Risk Indicators", [
            "The dashboard does not only display the model score; it also shows risk indicators. These indicators are generated from the transaction values submitted by the user. For example, if the amount is above a typical threshold, the system displays a large-amount warning. If BVN match status is false, the system displays an identity mismatch warning. If the receiver bank is a fintech wallet, the system identifies it as a possible cash-out channel.",
            "The purpose of the risk indicators is to make the result easier to understand. A fraud probability alone may not be meaningful to every user. By showing the reasons that contributed to suspicion, the system becomes more useful for administrators. This approach supports explainability at the interface level, even though the internal XGBoost model is more complex than a simple rule list (Nobel et al., 2024; Zhang et al., 2025).",
            "The indicators also help during testing. When the high-risk pattern button is used, the dashboard fills the form with suspicious values. The user can then compare the displayed risk indicators with the final fraud result. This makes it easier to demonstrate the behavior of the system during project defense or presentation.",
        ]),
        ("4.24 Implementation of Transaction History", [
            "Transaction history was implemented to allow users to review previous prediction results. Each time a transaction is scored, the backend stores the transaction ID, user ID, amount, sender bank, receiver bank, channel, BVN match status, fraud probability, fraud status, risk level, recommendation, and scoring time. This record is saved in the predictions table.",
            "The transaction history page retrieves stored records from the backend. It displays them in a table so that administrators can see previous scores and decisions. This feature is important because fraud detection systems should support monitoring and review. If a transaction is later confirmed as fraudulent or legitimate, the stored history can support investigation and future model improvement.",
            "In a more advanced version, the history table can be extended to include confirmation feedback, reviewer notes, and case status. This would allow the system to support a complete fraud investigation workflow, not just prediction.",
        ]),
        ("4.25 Implementation of Rules Management", [
            "The rules management feature allows administrators to define fraud policies from the dashboard. The default rules include high amount alert, critical amount block, late-night window, BVN mismatch block, and USSD high-risk review. These rules are stored in a JSON file, which makes them easy to read and update.",
            "The rules page displays active policies in a table. Administrators can create new rules by providing rule name, condition type, condition value, enforcement action, description, and enabled status. They can also delete existing rules. This feature was included because many financial institutions use a combination of machine learning and business rules.",
            "In the current prototype, rule management is available as an administrative feature, while full rule enforcement can be improved in future work. A production version should connect rules directly to the prediction flow so that important rules can override or adjust the model recommendation when necessary.",
        ]),
        ("4.26 Implementation of User Management", [
            "User management was implemented to control access to the system. The users table stores username, password hash, role, full name, email, account status, creation time, and last login. Administrators can create and delete users from the dashboard. This prevents unauthorized persons from accessing sensitive transaction history or changing fraud rules.",
            "The system uses roles to separate access levels. An admin can manage users and rules, while staff users can be limited to monitoring or scoring activities. Role-based access is important in financial systems because not every user should have the same privileges. For example, a staff member may need to view transaction results but should not be allowed to create administrator accounts.",
            "The login process creates a session token after successful authentication. This token is sent with protected requests so that the backend can confirm the identity and role of the user. Although this is suitable for a prototype, a production version should add stronger password hashing, session expiry, two-factor authentication, and audit logs (CBN, 2024; Digital Policy Alert, 2024).",
        ]),
        ("4.27 Picture Placement Guide for Chapter Four", [
            "The implementation chapter should include screenshots because pictures help examiners understand the developed system. The first screenshot should be placed after the explanation of authentication and should show the login page. This proves that the system has access control. The second screenshot should be placed after the frontend implementation section and should show the main dashboard with the transaction form and result panel.",
            "The third and fourth screenshots should show the Normal Pattern and High-Risk Pattern simulation buttons in use. These screenshots are useful because they show how the system can be tested without manually typing every value. The fifth screenshot should show a low-risk transaction result, while the sixth should show a high-risk or critical transaction result. These two pictures make the prediction behavior clearer.",
            "Additional screenshots should show the transaction history page, rules management page, and user management page. These pictures should be labelled as Figure 4.7, Figure 4.8, and Figure 4.9 respectively. Each screenshot should be inserted close to the paragraph that explains the feature. This makes the documentation easier to follow and gives the reader visual evidence of implementation.",
        ]),
        ("4.28 Detailed Testing Discussion", [
            "Testing was carried out to confirm that the system behaves correctly under different conditions. The first test focused on server startup. The backend must load without errors and must also load the saved model and preprocessing pipeline. If the model is missing, the system should not silently continue; it should report that prediction is unavailable.",
            "The second test focused on valid transaction submission. A complete transaction request should return a fraud probability, risk level, and recommendation. The third test focused on invalid input. If a required field is missing, the backend should return a validation error. This is important because the model should not receive incomplete or badly formatted data.",
            "Authentication was also tested. A valid username and password should create a session token, while an invalid login should be rejected. Protected pages such as transaction history, rules, and user management should require authentication. These tests help confirm that the system is not only functional but also controlled.",
        ]),
        ("4.29 Interpretation of Results", [
            "The results obtained from testing show that the system can separate normal transaction patterns from suspicious transaction patterns. Normal transactions with moderate amounts, verified BVN status, safer channels, and normal banking hours tend to produce lower fraud probabilities. Suspicious transactions with high amounts, BVN mismatch, USSD or Web channel, fintech receiver, and late-night time tend to produce higher fraud probabilities.",
            "The result should be interpreted as a decision-support output, not as a final legal judgment. A high score suggests that the transaction needs attention, but further review may still be required. This is important because fraud detection systems must balance fraud prevention with customer experience. Blocking every suspicious transaction without review can affect genuine users.",
            "The project therefore demonstrates a practical fraud detection workflow. It receives data, processes the data, produces a risk score, explains visible risk indicators, stores the result, and allows administrators to review the history. This makes it more complete than a model-only implementation.",
        ]),
        ("4.30 Suggested Screenshots and Captions", [
            "The following screenshots should be added to the final printed version of Chapter Four. Figure 4.1 should show the login page with the title and authentication form. Figure 4.2 should show the main dashboard after successful login. Figure 4.3 should show the transaction form filled with a normal transaction pattern. Figure 4.4 should show the transaction form filled with a high-risk pattern.",
            "Figure 4.5 should show the prediction result for a safe or low-risk transaction, while Figure 4.6 should show the prediction result for a high-risk or critical transaction. These two screenshots are very important because they prove that the system can produce different outputs depending on transaction behavior.",
            "Figure 4.7 should show the transaction history table. Figure 4.8 should show the rules management page, including active rules and the rule creation form. Figure 4.9 should show the user management page. These screenshots should be placed close to the explanations of the features so that the reader can easily connect each picture to the implementation discussion.",
        ]),
    ],
    "CHAPTER FIVE": [
        ("5.11 Overall Project Reflection", [
            "The development of Gojo Sentinel shows that fraud detection is both a technical and practical problem. The technical side involves data preparation, model training, API development, and database design. The practical side involves usability, interpretation, access control, and administrative decision-making. A fraud detection system that is accurate but difficult to use may not be effective in a real organization.",
            "One important lesson from the project is that data quality is central to machine learning. Since real banking data was not available, synthetic data was used. This allowed the system to be built and tested, but it also means that the model should not be considered ready for production use. In a real deployment, the model would need to be trained and validated with actual transaction records under strict privacy and regulatory controls.",
            "Another lesson is that machine learning works better when combined with human oversight. The system can produce risk scores quickly, but administrators still need to review cases, manage rules, and improve the system with feedback. This supports the idea that AI should assist decision-making rather than replace every human decision.",
        ]),
        ("5.12 Practical Implications of the System", [
            "If improved and deployed in a real environment, the system can help financial institutions reduce response time to suspicious transactions. A transaction can be scored immediately, and high-risk cases can be flagged for review. This can reduce the delay between fraud attempt and fraud response. It can also help staff focus on transactions that carry higher risk instead of reviewing all transactions manually.",
            "The dashboard can also support training and awareness. Staff can use the simulation feature to understand how different transaction features affect fraud risk. For example, they can compare a normal NIP transfer with a late-night USSD transfer involving BVN mismatch and a fintech receiver. This makes fraud patterns easier to understand.",
            "For academic use, the system provides a complete example of applied artificial intelligence. It connects machine learning theory with software engineering practice. Students can study the data generation process, preprocessing pipeline, model training, API design, frontend development, and testing approach.",
        ]),
        ("5.13 Ethical and Privacy Considerations", [
            "Fraud detection systems deal with sensitive financial and personal information. In real-world use, the system must protect customer data and follow privacy laws. Transaction records may contain account numbers, user identities, timestamps, and behavioral information. Such data must be stored securely, accessed only by authorized users, and used only for legitimate security purposes.",
            "Another ethical concern is fairness. If a model is trained on poor or biased data, it may wrongly classify some users as risky. This can affect customer experience and trust. For this reason, production fraud detection systems should be tested carefully across different customer groups and transaction types. They should also include a way for decisions to be reviewed.",
            "Transparency is also important. Users and administrators should understand that the model provides probability-based support, not absolute truth. The risk indicators included in Gojo Sentinel help improve transparency by showing visible reasons why a transaction may be considered risky.",
        ]),
        ("5.14 Future System Architecture", [
            "A future version of the system can be designed as a more advanced fraud monitoring platform. The backend can be connected to live transaction streams through secure banking APIs. A message queue can be added to handle large transaction volumes, while the prediction service can run separately from the dashboard. This would make the system more scalable.",
            "A feedback service can also be added. When administrators review a transaction, they can mark it as confirmed fraud or confirmed legitimate. These labels can be stored and used for periodic retraining. Over time, this would allow the model to learn from new fraud patterns and improve its performance.",
            "The system can also include a notification layer. High-risk transactions can trigger email alerts, SMS alerts, or dashboard notifications. This would make the system more useful for real-time monitoring. Audit logs should also be added so that every admin action is recorded.",
        ]),
        ("5.15 Final Conclusion", [
            "In conclusion, Gojo Sentinel demonstrates how an AI-based fraud detection system can be designed and implemented for Nigerian digital banking transactions. The system combines machine learning, backend services, database storage, authentication, rules management, and a user-friendly dashboard. It shows that fraud detection can be improved by analyzing transaction behavior instead of relying only on manual checks or fixed rules.",
            "The project achieved its main objectives by generating transaction data, training an XGBoost model, building a FastAPI backend, creating a frontend dashboard, adding transaction history, and supporting administrative functions. Although the system is not yet ready for real banking deployment, it provides a strong prototype and a good foundation for future research and development.",
            "With real data, stronger security, full rule enforcement, feedback-based retraining, and live API integration, the system can be developed into a more powerful fraud detection platform. The work therefore contributes both academically and practically to the study of artificial intelligence in financial security.",
        ]),
        ("5.16 Closing Note", [
            "The project confirms that fraud detection is not a single-step activity. It requires data, models, rules, user interfaces, security controls, and continuous improvement. Gojo Sentinel brings these parts together in one prototype. It gives a clear demonstration of how artificial intelligence can be applied to a problem that affects real people and financial institutions.",
            "The system also shows that local context matters. A fraud detection system for Nigeria should consider Nigerian transaction channels, local banks, fintech wallets, and BVN-related identity checks. By including these features, the project becomes more relevant to the environment it was designed for.",
            "Although there is room for improvement, the project provides a strong academic foundation. It can be defended as a practical system, improved as a research project, and extended as a real-world application if future work includes real data and stronger deployment controls.",
        ]),
    ],
}


def copy_paragraph(dst, src_p):
    text = src_p.text
    if not text.strip():
        dst.add_paragraph()
        return
    if text.strip() in {"CHAPTER TWO: LITERATURE REVIEW", "CHAPTER THREE", "CHAPTER FOUR", "CHAPTER FIVE"}:
        dst.add_page_break()
    style_name = src_p.style.name if src_p.style is not None else "Normal"
    if style_name not in dst.styles:
        style_name = "Normal"
    p = dst.add_paragraph(style=style_name)
    p.alignment = src_p.alignment
    run = p.add_run(text)
    run.bold = any(r.bold for r in src_p.runs) or style_name.startswith("Heading") or text.strip().startswith("CHAPTER")
    set_run_font(run)
    if style_name == "Normal" and text.strip() and not text.strip().startswith("[Insert"):
        p.paragraph_format.first_line_indent = Inches(0.5)


def copy_table(dst, src_table):
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    table = dst.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(src_table.rows):
        for c_idx, cell in enumerate(row.cells):
            target = table.cell(r_idx, c_idx)
            target.text = cell.text
            for p in target.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.bold = r_idx == 0
                    set_run_font(run)
    dst.add_paragraph()


def iter_blocks(doc):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def build():
    src = Document(SOURCE)
    dst = Document()
    # Remove the default empty paragraph when python-docx creates one.
    if dst.paragraphs:
        dst._body._element.remove(dst.paragraphs[0]._p)

    in_references = False
    inserted = set()
    current_chapter = None

    for block in iter_blocks(src):
        if hasattr(block, "text"):
            text = block.text.strip()
            if text == "REFERENCES":
                if current_chapter in EXPANSIONS and current_chapter not in inserted:
                    for heading, paragraphs in EXPANSIONS[current_chapter]:
                        add_heading(dst, heading, 2)
                        for para in paragraphs:
                            add_para(dst, para)
                    inserted.add(current_chapter)
                in_references = True
            if text in EXPANSIONS and current_chapter is not None and current_chapter not in inserted:
                for heading, paragraphs in EXPANSIONS[current_chapter]:
                    add_heading(dst, heading, 2)
                    for para in paragraphs:
                        add_para(dst, para)
                inserted.add(current_chapter)
            copy_paragraph(dst, block)
            if text in EXPANSIONS:
                current_chapter = text
        else:
            copy_table(dst, block)

    normalize(dst)
    dst.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
